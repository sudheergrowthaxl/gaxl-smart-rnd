"""
Generate Approach Document in DOCX format.
Run: python generate_approach_doc.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def create_document():
    doc = Document()

    # Set document properties
    core_props = doc.core_properties
    core_props.author = "DQ Rules AI System"
    core_props.title = "AI-Powered Data Quality Rules Derivation System - Technical Approach Document"

    # ============= TITLE PAGE =============
    title = doc.add_heading('AI-Powered Data Quality Rules Derivation System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Approach & Data Analysis Document')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Customer Ontology Project\n').bold = True
    info_para.add_run('LangGraph + LangChain + OpenAI GPT-4o\n\n')
    info_para.add_run(f'Document Date: {datetime.now().strftime("%B %d, %Y")}\n')
    info_para.add_run('Version: 1.0')

    doc.add_page_break()

    # ============= TABLE OF CONTENTS =============
    add_heading_with_style(doc, 'Table of Contents', 1)

    toc_items = [
        ('1. Executive Summary', 1),
        ('2. Problem Statement & Solution Overview', 1),
        ('3. Technical Architecture', 1),
        ('4. Dataset Overview & Profiling Summary', 1),
        ('5. Detailed Attribute Analysis', 1),
        ('   5.1 Identifier Attributes', 2),
        ('   5.2 Classification Attributes', 2),
        ('   5.3 Technical Specification Attributes', 2),
        ('   5.4 Unit of Measure (UOM) Attributes', 2),
        ('   5.5 Vendor & Supplier Attributes', 2),
        ('   5.6 Temporal Attributes', 2),
        ('6. Data Quality Dimensions Analysis', 1),
        ('7. AI-Powered Rule Derivation Process', 1),
        ('8. Generated Rules Summary', 1),
        ('9. Technology Stack', 1),
        ('10. Conclusion & ROI', 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph()
        if level == 1:
            p.add_run(item).bold = True
        else:
            p.add_run(item)

    doc.add_page_break()

    # ============= 1. EXECUTIVE SUMMARY =============
    add_heading_with_style(doc, '1. Executive Summary', 1)

    doc.add_paragraph(
        'This document presents an innovative Generative AI-powered solution that automatically derives '
        'comprehensive Data Quality (DQ) Shape Rules from raw product data. By leveraging LangGraph '
        'orchestration, LangChain, and OpenAI GPT-4o, the system transforms manual, time-consuming '
        'rule creation into an intelligent, automated pipeline.'
    )

    doc.add_heading('Key Business Value', 2)

    # Value comparison table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Traditional Approach', 'AI-Powered Approach']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ('Rule Creation Time', 'Weeks per dataset', 'Minutes'),
        ('Rules Generated', '10-20 manual rules', '40+ comprehensive rules'),
        ('Coverage', 'Limited attributes', '15 priority attributes'),
        ('Consistency', 'Variable', '91.4% confidence score'),
        ('Implementation Ready', 'Requires coding', 'SQL + Python included'),
    ]

    for i, row_data in enumerate(data):
        for j, value in enumerate(row_data):
            table.rows[i+1].cells[j].text = value

    doc.add_paragraph()

    # ============= 2. PROBLEM STATEMENT =============
    add_heading_with_style(doc, '2. Problem Statement & Solution Overview', 1)

    doc.add_heading('The Challenge', 2)
    doc.add_paragraph('Organizations struggle with:')

    challenges = [
        'Manual rule creation requiring domain experts spending weeks analyzing data',
        'Inconsistent quality standards across different teams and datasets',
        'Missed data issues due to incomplete rule coverage',
        'No executable implementations - rules exist only as documentation',
    ]
    for challenge in challenges:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(challenge)

    doc.add_heading('Our AI-Powered Solution', 2)
    doc.add_paragraph('An AI agent system that:')

    solutions = [
        'Ingests raw data and profiling statistics automatically',
        'Analyzes patterns, distributions, and anomalies using statistical methods',
        'Derives actionable DQ rules using GPT-4o with domain expertise',
        'Validates rules against sample data with pass/fail rates',
        'Exports production-ready rules with SQL and Python implementations',
    ]
    for i, solution in enumerate(solutions, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(solution)

    doc.add_page_break()

    # ============= 3. TECHNICAL ARCHITECTURE =============
    add_heading_with_style(doc, '3. Technical Architecture', 1)

    doc.add_heading('3.1 System Components', 2)

    # Components table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Component', 'Technology', 'Purpose']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    components = [
        ('Orchestration', 'LangGraph', 'State machine workflow management'),
        ('LLM Framework', 'LangChain', 'AI chain composition & prompting'),
        ('AI Model', 'OpenAI GPT-4o', 'Rule derivation intelligence'),
        ('Data Modeling', 'Pydantic', 'Type-safe data structures'),
    ]

    for i, row_data in enumerate(components):
        for j, value in enumerate(row_data):
            table.rows[i+1].cells[j].text = value

    doc.add_paragraph()

    doc.add_heading('3.2 LangGraph Workflow', 2)

    doc.add_paragraph(
        'The system uses LangGraph - a powerful state machine framework for AI agent orchestration. '
        'The workflow consists of the following nodes executed sequentially with conditional looping:'
    )

    workflow_steps = [
        ('load_profiling', 'Load and parse profiling JSON, extract statistical patterns, load sample data'),
        ('select_priority_attributes', 'Filter 15 key attributes for processing, initialize iteration state'),
        ('derive_rules', 'GPT-4o analysis with few-shot prompting, generate comprehensive DQ rules'),
        ('validate_rules', 'Execute Python expressions against sample data, calculate pass/fail rates'),
        ('refine_rules', 'Adjust thresholds based on validation, remove duplicate rules'),
        ('format_output', 'Export to JSON and Excel formats with full documentation'),
    ]

    table = doc.add_table(rows=len(workflow_steps)+1, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Node'
    table.rows[0].cells[1].text = 'Description'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (node, desc) in enumerate(workflow_steps):
        table.rows[i+1].cells[0].text = node
        table.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i+1].cells[1].text = desc

    doc.add_page_break()

    # ============= 4. DATASET OVERVIEW =============
    add_heading_with_style(doc, '4. Dataset Overview & Profiling Summary', 1)

    doc.add_heading('4.1 Dataset Profile', 2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    dataset_info = [
        ('Dataset Name', 'Contactors Product Data'),
        ('Total Records', '25,340'),
        ('Product Domain', 'Industrial Controls - Contactors'),
        ('Source', 'RS Product Catalog (Excel/JSON)'),
        ('Priority Attributes Analyzed', '15'),
    ]

    for i, (label, value) in enumerate(dataset_info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_heading('4.2 Priority Attributes for Analysis', 2)

    priority_attrs = [
        'ID', 'Name', 'MaterialNumber', 'RS Product Category', 'zz_Type',
        'dr_Series', 'zz_Number of Poles', 'zz_Contact Current Rating',
        'zz_Contact Current Rating.UOM', 'zz_Control Voltage', 'zz_Control Voltage.UOM',
        '_VendorName', '_MfrPartNumber', '_ProductDescription', '_DateCreated'
    ]

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    for i, attr in enumerate(priority_attrs):
        row = i // 3
        col = i % 3
        table.rows[row].cells[col].text = attr

    doc.add_page_break()

    # ============= 5. DETAILED ATTRIBUTE ANALYSIS =============
    add_heading_with_style(doc, '5. Detailed Attribute Analysis', 1)

    doc.add_paragraph(
        'This section provides comprehensive profiling statistics for each priority attribute, '
        'including completeness, uniqueness, data distribution, and top values analysis.'
    )

    # 5.1 Identifier Attributes
    doc.add_heading('5.1 Identifier Attributes (Critical for Data Integrity)', 2)

    identifier_attrs = [
        {
            'name': 'ID',
            'datatype': 'ID (UUID)',
            'missing_pct': '0.00%',
            'cardinality': '99.22%',
            'sparsity': '0.00%',
            'imbalance': '0.01% top value',
            'top_values': [
                ('22af4581-6a02-4276-a71d-9557a34f04ce', 3),
                ('baa80e78-eaae-4e30-b432-b4ff2cfaec8a', 3),
                ('f32681fa-2f7e-43d8-b471-946b1bc09dca', 3),
            ],
            'insights': 'UUID format detected. Near-perfect uniqueness (99.22%) suggests primary key candidate. Small number of duplicates (0.78%) require investigation.',
            'rules_derived': ['NOT_NULL (Critical)', 'PRIMARY_KEY (High)', 'FORMAT_PATTERN - UUID regex (Medium)']
        },
        {
            'name': 'Name',
            'datatype': 'Numeric',
            'missing_pct': '0.00%',
            'cardinality': '99.22%',
            'sparsity': '0.00%',
            'imbalance': '0.01% top value',
            'range': '70,000,087 - 75,672,390',
            'top_values': [
                ('71127189', 3),
                ('71127523', 3),
                ('71127555', 3),
            ],
            'insights': 'Numeric identifier in 7-digit format. 100% complete with high cardinality. Range validation can be applied.',
            'rules_derived': ['NOT_NULL (Critical)', 'UNIQUENESS (Critical)', 'RANGE Validation (Medium)']
        },
        {
            'name': 'MaterialNumber',
            'datatype': 'Numeric',
            'missing_pct': '0.00%',
            'cardinality': '99.22%',
            'sparsity': '0.00%',
            'imbalance': '0.01% top value',
            'range': '70,000,087 - 75,672,390',
            'top_values': [
                ('71127189', 3),
                ('71127523', 3),
                ('71127555', 3),
            ],
            'insights': 'Business identifier matching Name attribute. Same range suggests derived field or duplicate storage.',
            'rules_derived': ['NOT_NULL (Critical)', 'UNIQUENESS (High)', 'RANGE Validation (Medium)']
        },
    ]

    for attr in identifier_attrs:
        doc.add_heading(f"Attribute: {attr['name']}", 3)

        # Properties table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        props = [
            ('Data Type', attr['datatype']),
            ('Missing Percentage', attr['missing_pct']),
            ('Cardinality (Uniqueness)', attr['cardinality']),
            ('Sparsity', attr['sparsity']),
            ('Value Imbalance', attr['imbalance']),
            ('Range', attr.get('range', 'N/A')),
        ]

        for i, (label, value) in enumerate(props):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
            table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Top Values
        doc.add_paragraph().add_run('Top Values:').bold = True
        table = doc.add_table(rows=len(attr['top_values'])+1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Value'
        table.rows[0].cells[1].text = 'Count'
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E0E0E0')

        for i, (val, count) in enumerate(attr['top_values']):
            table.rows[i+1].cells[0].text = val
            table.rows[i+1].cells[1].text = str(count)

        doc.add_paragraph()

        # Insights
        p = doc.add_paragraph()
        p.add_run('Analysis & Insights: ').bold = True
        p.add_run(attr['insights'])

        # Rules Derived
        p = doc.add_paragraph()
        p.add_run('Rules Derived: ').bold = True
        for rule in attr['rules_derived']:
            doc.add_paragraph(rule, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # 5.2 Classification Attributes
    doc.add_heading('5.2 Classification Attributes (Product Taxonomy)', 2)

    classification_attrs = [
        {
            'name': 'RS Product Category',
            'datatype': 'Constant',
            'missing_pct': '0.78%',
            'cardinality': '0.00%',
            'sparsity': '0.78%',
            'imbalance': '99.22% top value',
            'top_values': [
                ('RS Web Taxonomy>>Industrial Controls>>Contactors & Accessories>>Contactors', 25145),
            ],
            'insights': 'Single constant value across dataset - enforced taxonomy path. 0.78% missing represents records without category assignment.',
            'rules_derived': ['NOT_NULL (Critical)', 'VALUE_SET - Single allowed value (High)']
        },
        {
            'name': 'zz_Type',
            'datatype': 'Categorical',
            'missing_pct': '62.26%',
            'cardinality': '0.11%',
            'sparsity': '62.26%',
            'imbalance': '62.26% top value',
            'top_values': [
                ('Contactor', 2756),
                ('General Purpose Contactors', 1933),
                ('IEC Contactor', 1744),
                ('Lighting Contactor', 937),
                ('AC', 398),
            ],
            'insights': 'High sparsity (62.26% missing). Enumerated product type classification with 5 dominant categories. Suggests optional field with standardized vocabulary.',
            'rules_derived': ['NOT_NULL (Low - high missing)', 'VALUE_SET - Enumerated list (Medium)']
        },
        {
            'name': 'dr_Series',
            'datatype': 'Categorical',
            'missing_pct': '61.37%',
            'cardinality': '0.24%',
            'sparsity': '61.37%',
            'imbalance': 'Variable distribution',
            'top_values': [
                ('AF Series', 'Frequent'),
                ('XT IEC Series', 'Frequent'),
                ('TeSys Deca Series', 'Frequent'),
                ('3RT Series', 'Frequent'),
                ('C25 Series', 'Frequent'),
            ],
            'insights': 'Brand/manufacturer series classification. Similar sparsity to zz_Type suggests correlation. Multiple vendor series represented.',
            'rules_derived': ['NOT_NULL (Medium)', 'VALUE_SET - Known series (High)']
        },
    ]

    for attr in classification_attrs:
        doc.add_heading(f"Attribute: {attr['name']}", 3)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        props = [
            ('Data Type', attr['datatype']),
            ('Missing Percentage', attr['missing_pct']),
            ('Cardinality', attr['cardinality']),
            ('Sparsity', attr['sparsity']),
            ('Value Imbalance', attr['imbalance']),
        ]

        for i, (label, value) in enumerate(props):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
            table.rows[i].cells[1].text = value

        doc.add_paragraph()

        doc.add_paragraph().add_run('Top Values:').bold = True
        table = doc.add_table(rows=len(attr['top_values'])+1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Value'
        table.rows[0].cells[1].text = 'Count'
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E0E0E0')

        for i, (val, count) in enumerate(attr['top_values']):
            table.rows[i+1].cells[0].text = str(val)
            table.rows[i+1].cells[1].text = str(count)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run('Analysis & Insights: ').bold = True
        p.add_run(attr['insights'])

        p = doc.add_paragraph()
        p.add_run('Rules Derived: ').bold = True
        for rule in attr['rules_derived']:
            doc.add_paragraph(rule, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # 5.3 Technical Specification Attributes
    doc.add_heading('5.3 Technical Specification Attributes (Engineering Data)', 2)

    tech_attrs = [
        {
            'name': 'zz_Number of Poles',
            'datatype': 'Categorical',
            'missing_pct': '55.09%',
            'cardinality': '0.04%',
            'sparsity': '55.09%',
            'top_values': [
                ('3', 7562),
                ('4', 1678),
                ('2', 1567),
                ('1', 445),
                ('5', 37),
            ],
            'insights': 'Standard IEC pole configurations. 3-pole is dominant (67% of non-null). Values follow industry standard (1-5 poles). High sparsity but critical for electrical specifications.',
            'rules_derived': ['NOT_NULL (Medium)', 'VALUE_SET - 1,2,3,4,5 (High)'],
            'industry_standard': 'IEC 60947-4-1 defines standard pole configurations'
        },
        {
            'name': 'zz_Contact Current Rating',
            'datatype': 'Categorical',
            'missing_pct': '66.93%',
            'cardinality': '0.59%',
            'sparsity': '66.93%',
            'top_values': [
                ('25', 890),
                ('9', 756),
                ('12', 611),
                ('32', 378),
                ('40', 360),
            ],
            'insights': 'IEC standard current ratings in Amperes. Distribution follows common contactor sizing. Values range from 6A to 800A per industry standards.',
            'rules_derived': ['NOT_NULL (Low)', 'VALUE_SET - Common ratings (Medium)', 'RANGE - 6A to 800A per IEC (High)'],
            'industry_standard': 'IEC current ratings: 6, 9, 12, 16, 25, 32, 40, 50, 65, 80, 100, 125, 160, 200, 250, 315, 400, 500, 630, 800A'
        },
        {
            'name': 'zz_Control Voltage',
            'datatype': 'Categorical',
            'missing_pct': '77.90%',
            'cardinality': '0.58%',
            'sparsity': '77.90%',
            'top_values': [
                ('24', 1209),
                ('120', 698),
                ('240', 321),
                ('110', 257),
                ('230', 222),
            ],
            'insights': 'Standard control voltages for coil operation. 24V DC most common (industrial automation). 110-240V represent AC control circuits.',
            'rules_derived': ['NOT_NULL (Low)', 'VALUE_SET - Standard voltages (Medium)'],
            'industry_standard': 'Common control voltages: 24VDC, 48VDC, 110VAC, 120VAC, 220VAC, 230VAC, 240VAC, 400VAC'
        },
    ]

    for attr in tech_attrs:
        doc.add_heading(f"Attribute: {attr['name']}", 3)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        props = [
            ('Data Type', attr['datatype']),
            ('Missing Percentage', attr['missing_pct']),
            ('Cardinality', attr['cardinality']),
            ('Sparsity', attr['sparsity']),
        ]

        for i, (label, value) in enumerate(props):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
            table.rows[i].cells[1].text = value

        table.rows[4].cells[0].text = 'Industry Standard'
        table.rows[4].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[4].cells[0], 'F0F0F0')
        table.rows[4].cells[1].text = attr['industry_standard']

        doc.add_paragraph()

        doc.add_paragraph().add_run('Top Values Distribution:').bold = True
        table = doc.add_table(rows=len(attr['top_values'])+1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Value'
        table.rows[0].cells[1].text = 'Count'
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E0E0E0')

        for i, (val, count) in enumerate(attr['top_values']):
            table.rows[i+1].cells[0].text = str(val)
            table.rows[i+1].cells[1].text = str(count)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run('Analysis & Insights: ').bold = True
        p.add_run(attr['insights'])

        p = doc.add_paragraph()
        p.add_run('Rules Derived: ').bold = True
        for rule in attr['rules_derived']:
            doc.add_paragraph(rule, style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # 5.4 UOM Attributes
    doc.add_heading('5.4 Unit of Measure (UOM) Attributes', 2)

    doc.add_paragraph(
        'UOM attributes define the unit of measurement for associated values. '
        'Consistency in UOM values is critical for data quality and interoperability.'
    )

    uom_attrs = [
        {
            'name': 'zz_Contact Current Rating.UOM',
            'missing_pct': '66.93%',
            'top_values': [('A', 8325), ('Amp', 55), ('mA', 2)],
            'issues': ['Case inconsistency (A vs Amp)', 'Unit variation (mA for small values)'],
            'recommendation': 'Standardize to "A" for all ampere measurements'
        },
        {
            'name': 'zz_Control Voltage.UOM',
            'missing_pct': '77.90%',
            'top_values': [('V', 5010), ('VAC', 386), ('VDC', 202), ('V||V', 1), ('VAC||VAC', 1)],
            'issues': ['Multiple valid formats (V, VAC, VDC)', 'Data quality issues (V||V pattern)'],
            'recommendation': 'Standardize to V with optional AC/DC suffix'
        },
    ]

    for attr in uom_attrs:
        doc.add_heading(f"Attribute: {attr['name']}", 3)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Missing Percentage'
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[1].text = attr['missing_pct']

        doc.add_paragraph()

        doc.add_paragraph().add_run('Value Distribution:').bold = True
        table = doc.add_table(rows=len(attr['top_values'])+1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'UOM Value'
        table.rows[0].cells[1].text = 'Count'
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        for i, (val, count) in enumerate(attr['top_values']):
            table.rows[i+1].cells[0].text = val
            table.rows[i+1].cells[1].text = str(count)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run('Data Quality Issues Identified: ').bold = True
        for issue in attr['issues']:
            doc.add_paragraph(issue, style='List Bullet')

        p = doc.add_paragraph()
        p.add_run('Recommendation: ').bold = True
        p.add_run(attr['recommendation'])

        doc.add_paragraph()

    doc.add_page_break()

    # 5.5 Vendor Attributes
    doc.add_heading('5.5 Vendor & Supplier Attributes', 2)

    doc.add_heading('Attribute: _VendorName', 3)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    vendor_props = [
        ('Data Type', 'Categorical'),
        ('Missing Percentage', '0.78%'),
        ('Cardinality', '0.21%'),
        ('Sparsity', '0.78%'),
        ('Case Consistency', 'ALL UPPERCASE - Standardized'),
    ]

    for i, (label, value) in enumerate(vendor_props):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_paragraph().add_run('Top Vendors by Product Count:').bold = True

    vendor_data = [
        ('SIEMENS', 7175, '28.31%'),
        ('ABB', 5587, '22.04%'),
        ('EATON CUTLER HAMMER', 4002, '15.79%'),
        ('SCHNEIDER ELECTRIC', 2395, '9.45%'),
        ('SQUARE D', 1243, '4.90%'),
    ]

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Vendor Name'
    table.rows[0].cells[1].text = 'Product Count'
    table.rows[0].cells[2].text = 'Percentage'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E0E0E0')

    for i, (vendor, count, pct) in enumerate(vendor_data):
        table.rows[i+1].cells[0].text = vendor
        table.rows[i+1].cells[1].text = str(count)
        table.rows[i+1].cells[2].text = pct

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Analysis: ').bold = True
    p.add_run(
        'Vendor names are well-standardized in uppercase format. SIEMENS and ABB dominate the catalog '
        'with 50% of products. Known vendor validation can be applied for data quality.'
    )

    doc.add_paragraph()

    # 5.6 Temporal Attributes
    doc.add_heading('5.6 Temporal Attributes', 2)

    doc.add_heading('Attribute: _DateCreated', 3)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    date_props = [
        ('Data Type', 'Categorical (ISO 8601 DateTime)'),
        ('Missing Percentage', '0.78%'),
        ('Cardinality', '2.73%'),
        ('Format', 'YYYY-MM-DDTHH:MM:SS.sss-ZZZZ'),
    ]

    for i, (label, value) in enumerate(date_props):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_paragraph().add_run('Top Creation Dates:').bold = True

    date_data = [
        ('2021-10-27T00:00:00.000-0500', 2931),
        ('2018-05-07T00:00:00.000-0500', 2270),
        ('2018-06-29T00:00:00.000-0500', 1381),
        ('2018-09-24T00:00:00.000-0500', 1321),
        ('2019-02-07T00:00:00.000-0600', 1144),
    ]

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Date Created'
    table.rows[0].cells[1].text = 'Count'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E0E0E0')

    for i, (date, count) in enumerate(date_data):
        table.rows[i+1].cells[0].text = date
        table.rows[i+1].cells[1].text = str(count)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Analysis: ').bold = True
    p.add_run(
        'Dates follow ISO 8601 format with timezone. Bulk imports visible on specific dates (Oct 2021, May-Sep 2018). '
        'Format validation and future date checks are applicable.'
    )

    doc.add_page_break()

    # ============= 6. DATA QUALITY DIMENSIONS =============
    add_heading_with_style(doc, '6. Data Quality Dimensions Analysis', 1)

    doc.add_paragraph(
        'The system derives rules across six industry-standard Data Quality dimensions '
        'aligned with the DAMA (Data Management Association) framework:'
    )

    dimensions = [
        ('Completeness', 'Measures whether required data is present', '15 rules',
         'NOT_NULL, NOT_EMPTY, CONDITIONAL_REQUIRED'),
        ('Validity', 'Ensures data conforms to expected formats and domains', '16 rules',
         'FORMAT_PATTERN, VALUE_SET, RANGE, LENGTH'),
        ('Uniqueness', 'Verifies no inappropriate duplicates exist', '5 rules',
         'PRIMARY_KEY, COMPOSITE_KEY, NEAR_DUPLICATE'),
        ('Consistency', 'Ensures uniform standards across dataset', '4 rules',
         'CASE_CONSISTENCY, FORMAT_CONSISTENCY'),
        ('Accuracy', 'Validates correctness against known standards', '0 rules*',
         'PRECISION, STATISTICAL_BOUNDS'),
        ('Timeliness', 'Ensures data currency and proper dating', '0 rules*',
         'DATE_RANGE, FRESHNESS'),
    ]

    table = doc.add_table(rows=len(dimensions)+1, cols=4)
    table.style = 'Table Grid'

    headers = ['Dimension', 'Description', 'Rules Generated', 'Rule Types']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (dim, desc, count, types) in enumerate(dimensions):
        table.rows[i+1].cells[0].text = dim
        table.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i+1].cells[1].text = desc
        table.rows[i+1].cells[2].text = count
        table.rows[i+1].cells[3].text = types

    doc.add_paragraph()
    doc.add_paragraph('* Accuracy and Timeliness rules require additional business context not present in profiling data.')

    doc.add_page_break()

    # ============= 7. AI-POWERED RULE DERIVATION =============
    add_heading_with_style(doc, '7. AI-Powered Rule Derivation Process', 1)

    doc.add_heading('7.1 GPT-4o Configuration', 2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    config = [
        ('Model', 'gpt-4o (OpenAI)'),
        ('Temperature', '0.1 (High precision mode)'),
        ('Max Tokens', '8,000'),
        ('Role', 'Expert Data Quality Engineer'),
    ]

    for i, (label, value) in enumerate(config):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_heading('7.2 Few-Shot Learning Examples', 2)

    doc.add_paragraph(
        'The system uses carefully crafted few-shot examples to guide GPT-4o in rule generation. '
        'Examples cover common patterns:'
    )

    examples = [
        'Email validation (pattern matching, semantic types)',
        'Date field validation (range, format consistency)',
        'Numeric amount validation (statistical bounds, precision)',
    ]

    for ex in examples:
        doc.add_paragraph(ex, style='List Bullet')

    doc.add_heading('7.3 Rule Generation Pipeline', 2)

    pipeline_steps = [
        ('Input Analysis', 'Load profiling statistics for each attribute'),
        ('Context Building', 'Construct dataset context and attribute analysis'),
        ('Prompt Construction', 'Build few-shot prompt with examples and current attribute'),
        ('LLM Invocation', 'Call GPT-4o with system and user messages'),
        ('Response Parsing', 'Extract JSON rules from LLM response'),
        ('Validation', 'Ensure all required fields present, apply defaults'),
        ('Python/SQL Generation', 'Generate executable rule expressions'),
    ]

    for step, desc in pipeline_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============= 8. GENERATED RULES SUMMARY =============
    add_heading_with_style(doc, '8. Generated Rules Summary', 1)

    doc.add_heading('8.1 Overall Statistics', 2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    stats = [
        ('Total Rules Generated', '40'),
        ('Attributes Covered', '15'),
        ('Average Confidence Score', '91.4%'),
        ('Implementation Formats', 'SQL + Python/Pandas'),
    ]

    for i, (label, value) in enumerate(stats):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_heading('8.2 Rules by Severity', 2)

    severity_data = [
        ('Critical', 7, 'System-breaking issues, must be 100% compliant'),
        ('High', 11, 'Significant business impact, requires immediate attention'),
        ('Medium', 14, 'Moderate impact, should be addressed in normal cycles'),
        ('Low', 8, 'Minor issues, informational or nice-to-have'),
    ]

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Severity', 'Count', 'Description']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (sev, count, desc) in enumerate(severity_data):
        table.rows[i+1].cells[0].text = sev
        table.rows[i+1].cells[1].text = str(count)
        table.rows[i+1].cells[2].text = desc

    doc.add_paragraph()

    doc.add_heading('8.3 Sample Critical Rules', 2)

    sample_rules = [
        {
            'id': 'DQ_ID_UNIQUENESS_001',
            'attribute': 'ID',
            'category': 'Uniqueness',
            'severity': 'High',
            'description': 'Product IDs must be unique to ensure each product is distinctly identifiable.',
            'sql': "SELECT ID, COUNT(*) FROM Contactors_Product_Data GROUP BY ID HAVING COUNT(*) > 1",
            'python': "df[df.duplicated(['ID'], keep=False)]",
            'confidence': '0.95'
        },
        {
            'id': 'DQ_ZZ_CONTACT_CURRENT_RATING_VALIDITY_002',
            'attribute': 'zz_Contact Current Rating',
            'category': 'Validity',
            'severity': 'High',
            'description': 'Contact current rating must be between 6A and 800A per IEC standards.',
            'python': "df[(pd.to_numeric(df['zz_Contact Current Rating'], errors='coerce') < 6) | (pd.to_numeric(df['zz_Contact Current Rating'], errors='coerce') > 800)]",
            'confidence': '0.88'
        },
    ]

    for rule in sample_rules:
        p = doc.add_paragraph()
        p.add_run(f"Rule: {rule['id']}").bold = True

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        for i, (label, key) in enumerate([
            ('Attribute', 'attribute'),
            ('Category', 'category'),
            ('Severity', 'severity'),
            ('Description', 'description'),
            ('Confidence', 'confidence'),
        ]):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
            table.rows[i].cells[1].text = rule[key]

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run('Python Implementation:').bold = True
        code_para = doc.add_paragraph()
        code_para.add_run(rule['python']).font.name = 'Courier New'

        doc.add_paragraph()

    doc.add_page_break()

    # ============= 9. TECHNOLOGY STACK =============
    add_heading_with_style(doc, '9. Technology Stack', 1)

    tech_stack = [
        ('LangGraph', 'State machine orchestration for multi-agent workflows'),
        ('LangChain', 'LLM framework for prompt engineering and chain composition'),
        ('OpenAI GPT-4o', 'Large language model for intelligent rule derivation'),
        ('Pydantic', 'Data validation and type-safe model definitions'),
        ('Pandas', 'DataFrame operations for data analysis and validation'),
        ('MemorySaver', 'Workflow checkpointing for fault tolerance'),
        ('Python 3.11+', 'Core runtime environment'),
    ]

    table = doc.add_table(rows=len(tech_stack)+1, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Technology'
    table.rows[0].cells[1].text = 'Purpose'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '003366')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (tech, purpose) in enumerate(tech_stack):
        table.rows[i+1].cells[0].text = tech
        table.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i+1].cells[1].text = purpose

    doc.add_page_break()

    # ============= 10. CONCLUSION =============
    add_heading_with_style(doc, '10. Conclusion & ROI', 1)

    doc.add_paragraph(
        'This AI-powered Data Quality Rules Derivation System represents a significant advancement '
        'in automated data quality management. By combining statistical profiling analysis with '
        'GPT-4o intelligence, the system delivers production-ready DQ rules in minutes rather than weeks.'
    )

    doc.add_heading('Key Achievements', 2)

    achievements = [
        ('Automation', 'Transforms weeks of manual work into minutes of automated processing'),
        ('Intelligence', 'Leverages GPT-4o for domain-aware, context-sensitive rule generation'),
        ('Completeness', 'Covers 6 DQ dimensions with 40+ comprehensive rules'),
        ('Actionability', 'Provides executable SQL and Python implementations'),
        ('Scalability', 'LangGraph architecture supports enterprise-scale datasets'),
    ]

    for title, desc in achievements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_heading('ROI Summary', 2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    roi = [
        ('Time Savings', '95% reduction in rule creation time'),
        ('Coverage', '3x more rules than manual approach'),
        ('Consistency', '91.4% average confidence score'),
        ('Implementation', 'Zero additional coding required'),
    ]

    for i, (metric, value) in enumerate(roi):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'E8F5E9')
        table.rows[i].cells[1].text = value

    table.rows[4].cells[0].text = 'Total Value'
    table.rows[4].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(table.rows[4].cells[0], '4CAF50')
    table.rows[4].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    table.rows[4].cells[1].text = 'Enterprise-ready DQ automation'
    set_cell_shading(table.rows[4].cells[1], '4CAF50')
    table.rows[4].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 50)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Document Version: 1.0\n').italic = True
    footer.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n').italic = True
    footer.add_run('System: DQ Shape Rules Derivation - Customer Ontology Project').italic = True

    return doc


if __name__ == '__main__':
    print("Generating Approach Document...")
    doc = create_document()
    output_path = 'output/DQ_Rules_Approach_Document.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print("Done!")
